import os
import io
import json
import logging
import asyncio
from typing import Optional, List

from fastapi import FastAPI, UploadFile, File, HTTPException, Header, Depends
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv
import google.generativeai as genai
from google.oauth2 import id_token
from google.auth.transport import requests as google_requests
import docx
from pypdf import PdfReader

# Load environment variables
load_dotenv()

# Configure Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

# CORS Setup
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify the frontend domain
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configuration
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    logger.warning("GOOGLE_API_KEY not set in environment variables.")

# Configure Gemini
genai.configure(api_key=GOOGLE_API_KEY)
# User has access to advanced models. Using Gemini 2.5 Flash for speed.
MODEL_NAME = "gemini-2.5-flash" 

SYSTEM_PROMPT = """
    Rewrite the text to be unique and human-like, but keep the length similar to the original so it fits in the document layout. 
    Do not add conversational filler.
    """

GENERATION_CONFIG = {
    "temperature": 0.7,
    "top_p": 0.8,
    "top_k": 40,
    "max_output_tokens": 8192,
}

# Concurrency Control
MAX_CONCURRENT_REQUESTS = 15
semaphore = asyncio.Semaphore(MAX_CONCURRENT_REQUESTS)

# --- Auth Helper ---
async def verify_google_token(authorization: Optional[str] = Header(None)):
    if not authorization:
        raise HTTPException(status_code=401, detail="Authorization header missing")
    
    try:
        scheme, token = authorization.split()
        if scheme.lower() != 'bearer':
            raise HTTPException(status_code=401, detail="Invalid authentication scheme")
        
        # Verify token
        # Note: In a real app, you should verify the AUDience (Client ID) too.
        # For this demo, we just check if it's a valid Google token.
        id_info = id_token.verify_oauth2_token(token, google_requests.Request())
        
        return id_info
    except ValueError as e:
        logger.error(f"Token verification failed: {e}")
        raise HTTPException(status_code=401, detail="Invalid token")
    except Exception as e:
        logger.error(f"Auth error: {e}")
        raise HTTPException(status_code=401, detail="Authentication failed")

# --- Document Processing ---
async def paraphrase_text(text: str) -> str:
    if not text.strip():
        return text
    
    async with semaphore:
        try:
            model = genai.GenerativeModel(
                model_name=MODEL_NAME,
                generation_config=GENERATION_CONFIG,
                system_instruction=SYSTEM_PROMPT
            )
            
            # Run in a thread pool to avoid blocking the event loop with synchronous API calls if the lib is sync
            # The google-generativeai lib's generate_content_async is preferred if available, 
            # but generate_content is synchronous.
            response = await model.generate_content_async(text)
            return response.text.strip()
        except Exception as e:
            logger.error(f"Gemini API Error for text '{text[:20]}...': {e}")
            # Fallback: return original text if paraphrasing fails to avoid breaking the doc
            return text

async def paraphrase_docx_inplace(file_bytes: bytes) -> io.BytesIO:
    doc = docx.Document(io.BytesIO(file_bytes))
    
    tasks = []
    elements = [] # Keep track of elements to update corresponding to tasks

    # 1. Iterate Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            tasks.append(paraphrase_text(para.text))
            elements.append(para)

    # 2. Iterate Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Recursively check paragraphs in cells (cells can contain paragraphs)
                for para in cell.paragraphs:
                    if para.text.strip():
                        tasks.append(paraphrase_text(para.text))
                        elements.append(para)

    logger.info(f"Found {len(tasks)} text elements to paraphrase.")

    # 3. Execute Concurrently
    if tasks:
        results = await asyncio.gather(*tasks)
        
        # 4. Update Document
        for element, new_text in zip(elements, results):
            # Preserving run formatting is hard if we replace the whole text.
            # Simple replacement:
            element.text = new_text
            
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# --- Legacy PDF Processing (Fallback) ---
def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def create_docx_from_text(text: str) -> io.BytesIO:
    doc = docx.Document()
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def chunk_text(text: str, chunk_size: int = 4000) -> List[str]:
    return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]

# --- Endpoints ---

@app.get("/")
def read_root():
    return {"message": "Undetectable Document Paraphraser API is running."}

@app.post("/paraphrase")
async def paraphrase_document(
    file: UploadFile = File(...),
    user_info: dict = Depends(verify_google_token)
):
    logger.info(f"Processing file: {file.filename} for user: {user_info.get('email')}")
    
    content = await file.read()
    
    if file.filename.endswith('.docx'):
        # New In-Place Logic
        output_io = await paraphrase_docx_inplace(content)
        
    elif file.filename.endswith('.pdf'):
        # Legacy Logic
        original_text = extract_text_from_pdf(content)
        if not original_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from file.")
            
        chunks = chunk_text(original_text, chunk_size=5000) 
        paraphrased_chunks = []
        for i, chunk in enumerate(chunks):
            processed = await paraphrase_text(chunk)
            paraphrased_chunks.append(processed)
        
        final_text = "\n".join(paraphrased_chunks)
        output_io = create_docx_from_text(final_text)
        
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Please upload .docx or .pdf")

    # Return File
    filename_base = os.path.splitext(file.filename)[0]
    output_filename = f"{filename_base}_humanized.docx"
    
    return StreamingResponse(
        output_io,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={output_filename}"}
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
